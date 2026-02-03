"""Document parser for extracting text content from various file formats.

Supported formats:
- .txt - Plain text
- .md - Markdown
- .docx - Microsoft Word 2007+ (requires python-docx)
- .pdf - PDF documents (requires PyPDF2 or pdfplumber)
"""

import os
from pathlib import Path
from typing import Optional


class DocumentParser:
    """Parse and extract text content from various document formats."""

    @staticmethod
    def extract_text(filepath: str) -> str:
        """Extract text content from a document file.

        Args:
            filepath: Path to the document file

        Returns:
            Extracted text content as string

        Raises:
            ValueError: If file format is not supported
            Exception: If file parsing fails
        """
        path = Path(filepath)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {filepath}")

        # Get file extension
        ext = path.suffix.lower()

        # Route to appropriate parser based on extension
        if ext in ['.txt', '.md']:
            return DocumentParser._extract_text_file(filepath)
        elif ext == '.docx':
            return DocumentParser._extract_docx(filepath)
        elif ext == '.pdf':
            return DocumentParser._extract_pdf(filepath)
        elif ext == '.doc':
            return DocumentParser._extract_doc(filepath)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    @staticmethod
    def _extract_text_file(filepath: str) -> str:
        """Extract text from .txt or .md files."""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(filepath, 'r', encoding='latin-1') as f:
                    return f.read()
            except Exception as e:
                raise Exception(f"Failed to read text file: {e}")

    @staticmethod
    def _extract_docx(filepath: str) -> str:
        """Extract text from .docx files using python-docx."""
        try:
            from docx import Document
        except ImportError:
            raise Exception(
                "python-docx library is required to read .docx files. "
                "Install it with: pip install python-docx"
            )

        try:
            doc = Document(filepath)
            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)

            return '\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Failed to parse .docx file: {e}")

    @staticmethod
    def _extract_pdf(filepath: str) -> str:
        """Extract text from .pdf files using PyPDF2 or pdfplumber."""
        # Try pdfplumber first (better quality)
        try:
            import pdfplumber
            return DocumentParser._extract_pdf_plumber(filepath, pdfplumber)
        except ImportError:
            pass

        # Fallback to PyPDF2
        try:
            import PyPDF2
            return DocumentParser._extract_pdf_pypdf2(filepath, PyPDF2)
        except ImportError:
            raise Exception(
                "Neither pdfplumber nor PyPDF2 is installed. "
                "Install one with: pip install pdfplumber or pip install PyPDF2"
            )

    @staticmethod
    def _extract_pdf_plumber(filepath: str, pdfplumber) -> str:
        """Extract text from PDF using pdfplumber."""
        text_content = []
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_content.append(text)
            return '\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Failed to parse PDF with pdfplumber: {e}")

    @staticmethod
    def _extract_pdf_pypdf2(filepath: str, PyPDF2) -> str:
        """Extract text from PDF using PyPDF2."""
        text_content = []
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    try:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
                    except Exception as e:
                        print(f"Warning: Failed to extract text from a page: {e}")
                        continue
            return '\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Failed to parse PDF with PyPDF2: {e}")

    @staticmethod
    def _extract_doc(filepath: str) -> str:
        """Extract text from .doc files (legacy Word format).

        Note: .doc files are harder to parse. We try multiple methods.
        """
        # Try python-docx2txt first
        try:
            import docx2txt
            return docx2txt.process(filepath)
        except ImportError:
            pass

        # Try textract
        try:
            import textract
            return textract.process(filepath).decode('utf-8')
        except ImportError:
            pass

        # Try antiword (Linux only, requires antiword installed)
        try:
            import subprocess
            result = subprocess.run(
                ['antiword', filepath],
                capture_output=True,
                text=True,
                timeout=30
            )
            if result.returncode == 0:
                return result.stdout
        except FileNotFoundError:
            pass
        except Exception as e:
            print(f"Warning: antiword failed: {e}")

        raise Exception(
            "Unable to parse .doc files. Install one of:\n"
            "- pip install docx2txt\n"
            "- pip install textract\n"
            "- apt-get install antiword (Linux only)"
        )

    @staticmethod
    def is_supported_format(filename: str) -> bool:
        """Check if file format is supported.

        Args:
            filename: Name of the file to check

        Returns:
            True if format is supported, False otherwise
        """
        ext = Path(filename).suffix.lower()
        return ext in ['.txt', '.md', '.docx', '.pdf', '.doc']

    @staticmethod
    def get_supported_extensions() -> list[str]:
        """Get list of supported file extensions.

        Returns:
            List of supported file extensions (with dots)
        """
        return ['.txt', '.md', '.docx', '.pdf', '.doc']
